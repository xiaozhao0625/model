from __future__ import annotations

import argparse
import re
from pathlib import Path


REPO_ROOT = Path(__file__).resolve().parents[3]
SOURCE_FILES = [
    "docs/deployment/p13_real_deploy_overview.md",
    "docs/deployment/p13_software_download_links.md",
    "docs/deployment/p13_m0_master_setup.md",
    "docs/deployment/p13_w1_pc_game_worker_setup.md",
    "docs/deployment/p13_w2_pc_app_web_worker_setup.md",
    "docs/deployment/p13_w3_android_worker_setup.md",
    "docs/deployment/p13_env_templates.md",
    "docs/deployment/p13_startup_order.md",
    "docs/deployment/p13_smoke_test_plan.md",
    "docs/deployment/p13_diagnostics_collection.md",
    "docs/deployment/p13_troubleshooting.md",
    "docs/deployment/p13_acceptance_checklist.md",
]
DEFAULT_OUTPUT = "docs/deployment/export/P13_四机真实部署与验收手册.docx"


def main() -> None:
    parser = argparse.ArgumentParser(description="Generate the P13 real deployment DOCX handbook.")
    parser.add_argument("--output", default=DEFAULT_OUTPUT)
    args = parser.parse_args()
    build_docx(Path(args.output))


def build_docx(output_path: Path) -> None:
    try:
        from docx import Document
        from docx.enum.section import WD_SECTION
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Inches, Pt, RGBColor
    except ImportError as exc:  # pragma: no cover - depends on doc tooling env
        raise SystemExit("python-docx is required to generate DOCX") from exc

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(11)
    for style_name, size, color in [
        ("Heading 1", 16, "2E74B5"),
        ("Heading 2", 13, "2E74B5"),
        ("Heading 3", 12, "1F4D78"),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(5)

    code_style = styles.add_style("P13 Code", 1)
    code_style.font.name = "Consolas"
    code_style.font.size = Pt(9)
    code_style.paragraph_format.left_indent = Inches(0.18)
    code_style.paragraph_format.space_after = Pt(3)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("P13 四机真实部署与验收手册")
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor.from_string("0B2545")
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("Master + Worker 四机手动部署、健康检查、smoke 验收与 diagnostics 收集").italic = True
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run("安全边界：不包含真实密码、不包含真实 DATABASE_URL、不下载或安装软件。")
    doc.add_page_break()

    doc.add_heading("章节索引", level=1)
    for index, source in enumerate(SOURCE_FILES, start=1):
        title_text = _extract_title(REPO_ROOT / source)
        doc.add_paragraph(f"{index}. {title_text}", style="List Number")
    doc.add_page_break()

    for source in SOURCE_FILES:
        _append_markdown(doc, REPO_ROOT / source)
        doc.add_section(WD_SECTION.NEW_PAGE)

    _set_footer(doc, "P13 四机真实部署与验收手册 | 不含真实凭据")
    output = REPO_ROOT / output_path
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)
    print(f"generated {output}")


def _extract_title(path: Path) -> str:
    for line in path.read_text(encoding="utf-8").splitlines():
        if line.startswith("# "):
            return line.removeprefix("# ").strip()
    return path.stem


def _append_markdown(doc, path: Path) -> None:
    lines = path.read_text(encoding="utf-8").splitlines()
    in_code = False
    table_buffer: list[str] = []
    for raw in lines:
        line = raw.rstrip()
        if line.startswith("```"):
            _flush_table(doc, table_buffer)
            table_buffer.clear()
            in_code = not in_code
            continue
        if in_code:
            if line:
                doc.add_paragraph(_sanitize(line), style="P13 Code")
            continue
        if line.startswith("|"):
            table_buffer.append(line)
            continue
        _flush_table(doc, table_buffer)
        table_buffer.clear()
        if not line.strip():
            continue
        if line.startswith("# "):
            doc.add_heading(_sanitize(line[2:].strip()), level=1)
        elif line.startswith("## "):
            doc.add_heading(_sanitize(line[3:].strip()), level=2)
        elif line.startswith("### "):
            doc.add_heading(_sanitize(line[4:].strip()), level=3)
        elif line.startswith("- [ ]"):
            doc.add_paragraph("□ " + _sanitize(line[5:].strip()), style="List Bullet")
        elif line.startswith("- "):
            doc.add_paragraph(_sanitize(line[2:].strip()), style="List Bullet")
        elif re.match(r"^\d+\.\s+", line):
            doc.add_paragraph(_sanitize(re.sub(r"^\d+\.\s+", "", line)), style="List Number")
        elif line.startswith("> "):
            para = doc.add_paragraph(_sanitize(line[2:].strip()))
            para.runs[0].italic = True
        else:
            doc.add_paragraph(_sanitize(line))
    _flush_table(doc, table_buffer)


def _flush_table(doc, rows: list[str]) -> None:
    if len(rows) < 2:
        for row in rows:
            doc.add_paragraph(_sanitize(row))
        return
    parsed = [[_sanitize(cell.strip()) for cell in row.strip("|").split("|")] for row in rows]
    if all(set(cell.replace(" ", "")) <= {"-", ":"} for cell in parsed[1]):
        parsed.pop(1)
    if not parsed:
        return
    table = doc.add_table(rows=len(parsed), cols=len(parsed[0]))
    table.style = "Table Grid"
    table.autofit = True
    for r_index, row in enumerate(parsed):
        for c_index, cell_text in enumerate(row[: len(parsed[0])]):
            cell = table.cell(r_index, c_index)
            cell.text = cell_text
            if r_index == 0:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True


def _set_footer(doc, text: str) -> None:
    for section in doc.sections:
        paragraph = section.footer.paragraphs[0]
        paragraph.text = text


def _sanitize(text: str) -> str:
    return text.replace("DATABASE_URL=postgresql+psycopg://screenshot_app:<password>", "DATABASE_URL=postgresql+psycopg://screenshot_app:<password>")


if __name__ == "__main__":
    main()
